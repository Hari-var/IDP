# import streamlit as st
# from PIL import Image
# import fitz  # PyMuPDF for PDF
# from docx import Document
# import extract_msg
# import email
# from email import policy
# import os
# from app import sql
# import pandas as pd

# st.set_page_config(layout="wide")
# st.title("File Viewer App")


# # Get file_path from query params
# # Get file_path from query params
# st.write('---')

# # Get file_path from query params
# query_params = st.query_params.to_dict()

# id = query_params.get("file_id")


# if not id:
#     st.error("No file selected or file does not exist.")
#     st.stop()

# file_name = os.path.basename(id)

# row= sql.get_details_by_id(id)
# # st.write(row)
# file_name = row[1]
# file_name_lower = file_name.lower()
# source=row[3]
# file_path=row[7]
# file_path = file_path.replace("\\", "/")  # Ensure path is in the correct format for Streamlit

# # col1, col2 = st.columns([1, 2])
# spacer1, col1, col2, spacer2 = st.columns([0.25, 1.75, 2, 0.5])


# with col2:
#     st.subheader("📄 File Preview")
#     # Image files
#     if file_name_lower.endswith(("jpg", "jpeg", "png", "tiff")):
#         img = Image.open(file_path)
#         st.image(img, caption=file_name, use_container_width=True)

#     # PDF files
#     elif file_name_lower.endswith("pdf"):
#         with open(file_path, "rb") as f:
#             pdf = fitz.open(stream=f.read(), filetype="pdf")
#         for page_number in range(len(pdf)):
#             page = pdf[page_number]
#             pix = page.get_pixmap()
#             img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
#             st.image(img, caption=f"Page {page_number + 1}")

#     # DOCX files
#     elif file_name_lower.endswith("docx"):
#         doc = Document(file_path)
#         st.markdown("### Document Content:")
#         for para in doc.paragraphs:
#             st.write(para.text)

#     # EML files
#     elif file_name_lower.endswith("eml"):
#         with open(file_path, "rb") as f:
#             eml_content = f.read()
#         msg = email.message_from_bytes(eml_content, policy=policy.default)
#         st.markdown("### Email Details:")
#         st.write("Subject:", msg["subject"])
#         st.write("From:", msg["from"])
#         st.write("To:", msg["to"])
#         st.markdown("### Body:")
#         if msg.is_multipart():
#             for part in msg.walk():
#                 if part.get_content_type() == "text/plain":
#                     st.text(part.get_payload(decode=True).decode(errors="ignore"))
#         else:
#             st.text(msg.get_payload(decode=True).decode(errors="ignore"))

#     # MSG files
#     elif file_name_lower.endswith("msg"):
#         msg = extract_msg.Message(file_path)
#         st.markdown("### Outlook Message:")
#         st.write("Subject:", msg.subject)
#         st.write("From:", msg.sender)
#         st.write("To:", msg.to)
#         st.markdown("### Body:")
#         st.text(msg.body)

#     elif file_name_lower.endswith("xls"):
#         df= pd.read_excel(file_path)
#         st.dataframe(df, use_container_width=True)
#     else:
#         st.warning("Unsupported file type.")

# with col1:
#     st.subheader("📝 File Details Form")
#     st.text_input("File Name", value=file_name,disabled=True)
#     prediction = st.text_input("Prediction", value=row[4], placeholder="Enter predicted document type")
#     summary = st.text_area("Summary", value=row[6], placeholder="Enter summary of the document", height=450)
#     Source = st.selectbox(
#     "Source",
#     ["Email", "API", "Manual Upload"],
#     index=["Email", "API", "Manual Upload"].index(source) if source in ["Email", "API", "Manual Upload"] else 0
# )
#     col1, col2 = st.columns([1, 2])
#     with col1:
#         if st.button("Submit",use_container_width=True):
            
#             sql.update_document_by_id(id, prediction, summary, Source)
#             st.success("Form submitted!")
            
#             st.write("📌 **Document is Updated:**")
#             # st.write(f"**File Name**: {file_name}")
#             # st.write(f"**Prediction**: {prediction}")
#             # st.write(f"**Summary**: {summary}")
#     with col2 :
#         if st.button("Delete Document",use_container_width=True):
#             sql.delete_document_by_id(id)
#             st.success("Document deleted successfully!")
#             st.write("Redirecting to the home page...")
#             st.rerun()
    
#     if st.button("Back to Home",use_container_width=True):
#             st.write("Redirecting to the home page...")
            
        
#             st.switch_page("pages/dashboard.py")


import streamlit as st
from PIL import Image
import fitz  # PyMuPDF for PDF
from docx import Document
import extract_msg
import email
from email import policy
import os
import time
# from app import sql
import pandas as pd
import requests

st.set_page_config(layout="wide")
st.title("File Viewer App")


# Get file_path from query params
# Get file_path from query params
st.write('---')

# Get file_path from query params
query_params = st.query_params.to_dict()

id = query_params.get("file_id")


if not id:
    st.error("No file selected or file does not exist.")
    st.stop()

file_name = os.path.basename(id)
#API call to get the file details by id
response=requests.get(f"http://localhost:8000/get_details_by_id/{id}")
# st.write(response)
if response.status_code == 200:
    row = response.json()
else:
    st.error("Failed to retrieve file details.")
    st.stop()
# st.dataframe(row)
file_name = row["document_name"]
file_name_lower = file_name.lower()
source=row["source"]
file_path=row["file_path"]
file_path = file_path.replace("\\", "/")  # Ensure path is in the correct format for Streamlit

# col1, col2 = st.columns([1, 2])
spacer1, col1, col2, spacer2 = st.columns([0.25, 1.75, 2, 0.5])


with col2:
    st.subheader("📄 File Preview")
    # Image files
    if file_name_lower.endswith(("jpg", "jpeg", "png", "tiff")):
        img = Image.open(file_path)
        st.image(img, caption=file_name, use_container_width=True)

    # PDF files
    elif file_name_lower.endswith("pdf"):
        with open(file_path, "rb") as f:
            pdf = fitz.open(stream=f.read(), filetype="pdf")
        for page_number in range(len(pdf)):
            page = pdf[page_number]
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            st.image(img, caption=f"Page {page_number + 1}")

    # DOCX files
    elif file_name_lower.endswith("docx"):
        doc = Document(file_path)
        st.markdown("### Document Content:")
        for para in doc.paragraphs:
            st.write(para.text)

    # EML files
    elif file_name_lower.endswith("eml"):
        with open(file_path, "rb") as f:
            eml_content = f.read()
        msg = email.message_from_bytes(eml_content, policy=policy.default)
        st.markdown("### Email Details:")
        st.write("Subject:", msg["subject"])
        st.write("From:", msg["from"])
        st.write("To:", msg["to"])
        st.markdown("### Body:")
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    st.text(part.get_payload(decode=True).decode(errors="ignore"))
        else:
            st.text(msg.get_payload(decode=True).decode(errors="ignore"))

    # MSG files
    elif file_name_lower.endswith("msg"):
        msg = extract_msg.Message(file_path)
        st.markdown("### Outlook Message:")
        st.write("Subject:", msg.subject)
        st.write("From:", msg.sender)
        st.write("To:", msg.to)
        st.markdown("### Body:")
        st.text(msg.body)
        if msg.attachments:
            st.markdown("### Attachments:")
            for attachment in msg.attachments:
                st.write(attachment.longFilename or attachment.shortFilename)

    elif file_name_lower.endswith("xls"):
        df= pd.read_excel(file_path)
        st.dataframe(df, use_container_width=True)
    else:
        st.warning("Unsupported file type.")

with col1:
    st.subheader("📝 File Details Form")
    st.text_input("File Name", value=file_name,disabled=True)
    prediction = st.text_input("Prediction", value=row['doc_type_predicted'], placeholder="Enter predicted document type")
    summary = st.text_area("Summary", value=row['summary'], placeholder="Enter summary of the document", height=450)
    Source = st.selectbox(
    "Source",
    ["Email", "API", "Manual Upload"],
    index=["Email", "API", "Manual Upload"].index(source) if source in ["Email", "API", "Manual Upload"] else 0
)
    col1, col2 = st.columns([1, 2])
    with col1:
        if st.button("Submit",use_container_width=True):
            #API call to update the document details
            response = requests.put(
                f"http://localhost:8000/update_document_by_id/{id}",
                json={
                    "file_name": file_name,
                    "prediction": prediction,
                    "summary": summary,
                    "source": Source
                }
            )
            if response.status_code == 200:
                st.success("Form submitted!")
            else:
                st.error("Failed to update document.")

            st.write("📌 **Document is Updated:**")
            # st.write(f"**File Name**: {file_name}")
            # st.write(f"**Prediction**: {prediction}")
            # st.write(f"**Summary**: {summary}")
    with col2 :
        if st.button("Delete Document",use_container_width=True):
            #API call to delete the document by id
            response = requests.delete(f"http://localhost:8000/delete_document_by_id/{id}")
            if response.status_code == 200:
                st.success("Document deleted successfully!")
                st.write("Redirecting to the home page...")
                time.sleep(2)
                st.switch_page("pages/dashboard.py")
            else:
                st.error("Failed to delete document.")
            
            
    
    if st.button("Back to Home",use_container_width=True):
            st.write("Redirecting to the home page...")
            st.switch_page("pages/dashboard.py")